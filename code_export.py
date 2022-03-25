import os
import pathlib
from docx import Document

#----Settings----

#input repo for export
directory_in_str = "/home/user/project"

#extensions for export
#extensions = list(['.html', '.css', '.js'])
extensions = list(['.java'])

#----------------

#path to files
output_file = 'test.temp'
doc_file = 'res.docx'

if os.path.isdir(directory_in_str):

    print("Directory: ", directory_in_str)
    print("Result doc file: ", doc_file)

    #check old files with the same names
    if os.path.isfile(output_file):
        os.remove(output_file)
    if os.path.isfile(doc_file):
        os.remove(doc_file)

    directory = os.fsencode(directory_in_str)
    file_counter = 0 # it needs for result check

    # walk throgh subdirectories in such directory
    for subdir, dirs, files in os.walk(directory):
        for file in files:
            filename = os.fsdecode(file)
            subdir_in_str = os.fsdecode(subdir)
            full_path = os.path.join(subdir_in_str, filename)
            ext = str(pathlib.Path(filename).suffix)
            if ext in extensions: # for only needed extensions
                with open(output_file,'a', encoding="utf-8") as fout:
                    with open(full_path,'r', encoding="utf-8") as f:
                        # print file name from directory
                        print('**' + str(full_path).replace(directory_in_str, ''), file=fout)
                        for line in f: # print every line
                            print(line.rstrip(), file=fout)
                        print(file=fout) #empty line to separate
                file_counter += 1

    print()
    print("Exported files to temporary file: ", file_counter)

    with open(output_file,'r', encoding="utf-8") as fout:
        document = Document()
        for line in fout:
            strip_line = line.rstrip()
            if strip_line:
                if strip_line[:2].count('*') == 2: # file name line has '**' symbols at beginning
                    document.add_paragraph().add_run(strip_line[2:]).bold = True # do it bold
                else:
                    document.add_paragraph(strip_line)
            else:
                document.add_paragraph()
        document.save(doc_file)
        print()
        print("Successfully export to word file! " + str(os.path.abspath(doc_file)))
    
    # remove temporary file
    if os.path.isfile(output_file):
        os.remove(output_file)
else:
    print("No such directory!")
